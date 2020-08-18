# -*- coding: UTF-8 -*-
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.enum.text import WD_COLOR_INDEX
from docx.enum.text import WD_LINE_SPACING
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK
from docx.shared import RGBColor
from docx.shared import Pt
from datetime import date
import re
import string
import requests
from bs4 import BeautifulSoup
from unicodedata import normalize
import pyodbc

from BancoNomes import Deputados,Comissoes,NomesIBGE,TMesAno

class conecta_banco_dados:
    def __init__(self):
        # Criando conexão.
        self.con = pyodbc.connect(r'DRIVER={Microsoft Access Driver (*.mdb)};'
                                  r'DBQ=/home/pi/projetos/ALERJ/ModuloJorge/ALERJ_NOTES.mdb;')
        # Criando cursor.
        self.cur = self.con.cursor()

def carrega_dominio_comissoes():
    #dict_nomes_comissoes ={}
    #chave = ''
    #banco = conecta_banco_dados()
    #sql = 'SELECT nome_comissao, link_www3, link_alerjln1 FROM nomes_comissoes'
    #lista_achou = banco.cur.execute(sql).fetchall()
    #for item in lista_achou:
    #    lista_atributos=[]
    #    chave= item[0].upper()
    #    chave = remove_acentos(chave)
    #    chave = re.sub(',', '', chave)
    #    lista_atributos.append(item[0])
    #    lista_atributos.append(item[1])
    #    lista_atributos.append(item[2])
    #    dict_nomes_comissoes.update({chave:lista_atributos})
    dict_nomes_comissoes ={}
    chave = ''
    #banco = conecta_banco_dados()
    #sql = 'SELECT nome_comissao, link_www3, link_alerjln1 FROM nomes_comissoes'
    #lista_achou = banco.cur.execute(sql).fetchall()
    lista_achou = Comissoes.comissoes
    for item in lista_achou:
        lista_atributos=[]
        chave= lista_achou[item]["nome"].upper()
        chave = remove_acentos(chave)
        chave = re.sub(',', '', chave)
        lista_atributos.append(lista_achou[item]["nome"])
        lista_atributos.append(lista_achou[item]["link_www3"])
        lista_atributos.append(lista_achou[item]["link_www3"])
        dict_nomes_comissoes.update({chave:lista_atributos})
    return(dict_nomes_comissoes)

def verifica_nome_sexo_deputado(texto):
    #banco = conecta_banco_dados()
    #texto_sem_acento = remove_acentos(texto)
    #lista_nome=texto_sem_acento.split(' ')
    #achou = 'Não'
    #sexo = 'M'
    #sexo_calculo = 1
    #sql = 'SELECT nome_deputado FROM nomes_deputados where nome_deputado = ?'
    #params = (texto)
    #lista_achou = banco.cur.execute(sql, params).fetchall()
    #if len(lista_achou) > 0:
    #   achou = "Sim"
    #   sql = 'SELECT classificacao FROM nomes_IBGE where primeiro_nome = ?'
    #   params = (lista_nome[0])
    #   lista_achou = banco.cur.execute(sql, params).fetchall()
    #   if len(lista_achou) > 0:
    #      sexo = lista_achou[0][0]
    #      if sexo == 'F':
    #         sexo_calculo = 0

    texto_sem_acento = remove_acentos(texto)
    lista_nome=texto_sem_acento.split(' ')
    achou = 'Nao'
    sexo = 'M'
    sexo_calculo = 1
    lista_achou = Deputados.localizarDeputado(texto)
    if len(lista_achou) > 0:
       achou = "Sim"
       params = (lista_nome[0])
       lista_achou = NomesIBGE.localizarNome(params)
       if len(lista_achou) > 0:
          sexo = lista_achou["sexo"]
          if sexo == 'F':
             sexo_calculo = 0
    return(achou,sexo_calculo)


def remove_acentos(texto):
    return normalize('NFKD', texto).encode('ASCII', 'ignore').decode('ASCII')


def gera_pauta(nome_doc_pauta,id_processamento,
               dict_sessao,
               dict_projeto_sessao,
               dict_projeto,
               dict_chave_projeto,
               dict_autores,
               dict_chave_ementa,
               dict_parecer,
               dict_novo_parecer,
               dict_relator,
               dict_comissao_link,
               dict_chave_parecer,
               dict_chave_novo_parecer,
               dict_chave_parecer_pendente,
               dict_tipo_projeto,
               dict_cores):
    lista_aux = []
    dict_aux = {}
    texto_chave = ''
    pos_inicial = 0
    
    lista_arquivos = []
    # -----------------------------------------Configuração do documento Web
    documento_web = Document()
    style = documento_web.styles['Normal']
    secao_atual = documento_web.sections[-1]
    font = style.font
    font.name = 'Ecofont Vera Sans'
    font.size = Pt(12)
    secao_atual.left_margin = Cm(1.5)
    secao_atual.right_margin = Cm(1.5)
    secao_atual.orientation = WD_ORIENT.PORTRAIT
    secao_atual.top_margin = Cm(1.5)
    secao_atual.bottom_margin = Cm(1.5)
    secao_atual.page_width = Cm(21.0)
    secao_atual.page_height = Cm(29.7)
    secao_atual.header_distance = Cm(1.25)
    secao_atual.footer_distance = Cm(1.25)

    lista_documentos =[]
    cont_sessao = 0
    for sessao in dict_sessao:
        lista_documentos.append(Document()) 
    lista_projetos_sessao = []
    lista_atributos_sessao = []
    cont_sessao = 0
    for sessao in dict_sessao:
        cont_sessao = cont_sessao + 1
        if cont_sessao > 1:
           p = documento_web.add_paragraph()
           run_web = p.add_run()
           run_web.add_break(WD_BREAK.PAGE)

        # -------------------------------------------------Quebra por Sessão (pauta)
        #
        # -------------------------------------------------Data da Ordem do Dia
        lista_atributos_sessao = dict_sessao[sessao]

        # -------------------------------------------------Documento Web
        p = documento_web.add_paragraph()
        p_formato = p.paragraph_format
        p_formato.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_formato.line_spacing = 1.0
        p_formato.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p_formato.space_after = 0
        p_formato.space_before = 0
        run_web = p.add_run(lista_atributos_sessao[1])
        run_web.font.size = Pt(18)
        run_web.bold = True
        #

        # -------------------------------------------------Documento Notes
        p_notes = lista_documentos[cont_sessao-1].add_paragraph()
        p_notes_formato = p_notes.paragraph_format
        p_notes_formato.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_notes_formato.line_spacing = 1.0
        p_notes_formato.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p_notes_formato.space_after = 0
        p_notes_formato.space_before = 0
        run_notes = p_notes.add_run(lista_atributos_sessao[1])
        run_notes.font.size = Pt(18)
        run_notes.bold = True
        #

        # ------------------------------------------ Dia da semana
        # -------------------------------------------------Documento Web
        p = documento_web.add_paragraph()
        p_formato = p.paragraph_format
        p_formato.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_formato.line_spacing = 1.0
        p_formato.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p_formato.space_after = 0
        run_web = p.add_run(lista_atributos_sessao[2].upper())
        run_web.font.size = Pt(10)
        run_web.bold = True
        run_web.font.color.rgb = RGBColor(dict_cores['CINZA'][0], dict_cores['CINZA'][1], dict_cores['CINZA'][2])

        # -------------------------------------------------Documento Notes
        p_notes = lista_documentos[cont_sessao-1].add_paragraph()
        p_notes_formato = p_notes.paragraph_format
        p_notes_formato.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_notes_formato.line_spacing = 1.0
        p_notes_formato.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p_notes_formato.space_after = 0
        run_notes = p_notes.add_run(lista_atributos_sessao[2].upper())
        run_notes.font.size = Pt(10)
        run_notes.bold = True
        run_notes.font.color.rgb = RGBColor(dict_cores['CINZA'][0], dict_cores['CINZA'][1], dict_cores['CINZA'][2])

        #
        # ------------------------------------------ Título da Sessão (Ordinária ou Extra)
        # -------------------------------------------------Documento Web
        p = documento_web.add_paragraph()
        p_formato = p.paragraph_format
        p_formato.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_formato.line_spacing = 1.0
        p_formato.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p_formato.space_after = 0
        run_web = p.add_run(lista_atributos_sessao[0])
        run_web.font.size = Pt(18)
        run_web.bold = True
        #
        # -------------------------------------------------Documento Notes
        p_notes = lista_documentos[cont_sessao-1].add_paragraph()
        p_notes_formato = p_notes.paragraph_format
        p_notes_formato.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_notes_formato.line_spacing = 1.0
        p_notes_formato.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p_notes_formato.space_after = 0
        run_notes = p_notes.add_run(lista_atributos_sessao[0])
        run_notes.font.size = Pt(18)
        run_notes.bold = True
        #

        # ------------------------------------------ Hora da Sessão (Ordinária ou Extra)
        # -------------------------------------------------Documento Web
        p = documento_web.add_paragraph()
        p_formato = p.paragraph_format
        p_formato.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_formato.line_spacing = 1.0
        p_formato.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p_formato.space_after = 0
        run_web = p.add_run(lista_atributos_sessao[3])
        run_web.font.size = Pt(18)
        run_web.bold = True
        #
        # -------------------------------------------------Documento Notes
        p_notes = lista_documentos[cont_sessao-1].add_paragraph()
        p_notes_formato = p_notes.paragraph_format
        p_notes_formato.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_notes_formato.line_spacing = 1.0
        p_notes_formato.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p_notes_formato.space_after = 0
        run_notes = p_notes.add_run(lista_atributos_sessao[3])
        run_notes.font.size = Pt(18)
        run_notes.bold = True
        #
        # --------------------------------------------Adiciona um parágrafo em branco
        # -------------------------------------------------Documento Web
        p = documento_web.add_paragraph()
        #
        # -------------------------------------------------Documento Notes
        p = lista_documentos[cont_sessao-1].add_paragraph()
        #
        texto_tramitacao_ant = ''
        texto_inclusao_ri_ant = ''
        #
        # --------------------------------------------------Lista proposições da Sessão
        lista_projetos_sessao = dict_projeto_sessao[sessao]  # Carregando a lista de projetos da sessão
        # texto_projeto, texto_ordinal_projeto, texto_inclusao_ri, lista_situacao_projeto, texto_autores, texto_ementa, flag_veto_parcial, flag_veto_total_flag_redacao_vencido, flag_redacao_final
        for item in lista_projetos_sessao:
            p = documento_web.add_paragraph()
            p_formato = p.paragraph_format
            p_formato.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p_formato.line_spacing = 1.0
            p_formato.line_spacing_rule = WD_LINE_SPACING.SINGLE
            p_formato.space_after = 0
            p_formato.space_before = 0
            #
            # ---------------------------------------------- Quebra por tipo de inclusão por regimento interno
            #
            texto_projeto = dict_projeto[item][0]
            texto_autores = dict_projeto[item][4]
            texto_autores = texto_autores.title()  # Tratar a preposição "Do" para ficar toda em caixa-baixa
            texto_ementa = dict_projeto[item][5]
            texto_inclusao_ri = dict_projeto[item][2]
            texto_ordinal_projeto = dict_projeto[item][1] + ' '
            if texto_inclusao_ri != texto_inclusao_ri_ant:
                p = documento_web.add_paragraph()
                p_formato = p.paragraph_format
                p_formato.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p_formato.line_spacing = 1.0
                p_formato.line_spacing_rule = WD_LINE_SPACING.SINGLE
                p_formato.space_after = 0
                p_formato.space_before = 0
                p = documento_web.add_paragraph()
                p_formato = p.paragraph_format
                p_formato.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p_formato.line_spacing = 1.0
                p_formato.line_spacing_rule = WD_LINE_SPACING.SINGLE
                p_formato.space_after = 0
                p_formato.space_before = 0
                run = p.add_run(dict_projeto[item][2])
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                #
                # --------------------------------------------Adiciona dois parágrafos em branco
                texto_inclusao_ri_ant = texto_inclusao_ri
            texto_tramitacao_1 = ''
            texto_tramitacao_2 = ''
            lista_situacao_projeto = dict_projeto[item][3]
            texto_tramitacao_1 = lista_situacao_projeto[0]
            #
            # ---------------------------------------------- Quebra por tipo de tramitação e regime
            #
            if len(lista_situacao_projeto) == 2:  # Prever mais de 2 linhas de situação de tramitação (laço for..)
                texto_tramitacao_2 = lista_situacao_projeto[1]
            if texto_tramitacao_1 + texto_tramitacao_2 != texto_tramitacao_ant:
                #
                # --------------------------------------------Adiciona grupo de Regime e Tramitação de uma ou mais proposições
                #
                # --------------------------------------------Primeira linha
                p = documento_web.add_paragraph()
                p_formato = p.paragraph_format
                p_formato.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p_formato.line_spacing = 1.0
                p_formato.line_spacing_rule = WD_LINE_SPACING.SINGLE
                p_formato.space_after = 0
                p_formato.space_before = 0
                p = documento_web.add_paragraph()
                run = p.add_run(texto_tramitacao_1)
                p_formato = p.paragraph_format
                p_formato.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_formato.line_spacing = 1.0
                p_formato.line_spacing_rule = WD_LINE_SPACING.SINGLE
                p_formato.space_after = 0
                p_formato.space_before = 0
                run.font.size = Pt(12)
                run.bold = True
                #
                # --------------------------------------------Segunda linha (pular uma linha e manter o parágrafo junto)
                p = documento_web.add_paragraph()
                run = p.add_run(texto_tramitacao_2)
                p_formato = p.paragraph_format
                p_formato.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_formato.line_spacing = 1.0
                p_formato.line_spacing_rule = WD_LINE_SPACING.SINGLE
                p_formato.space_after = 0
                p_formato.space_before = 0
                run.font.size = Pt(12)
                run.bold = True
                texto_tramitacao_ant = texto_tramitacao_1 + texto_tramitacao_2
                # --------------------------------------------Adiciona um parágrafo em branco
                p = documento_web.add_paragraph()
                p_formato = p.paragraph_format
                p_formato.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p_formato.line_spacing = 1.0
                p_formato.line_spacing_rule = WD_LINE_SPACING.SINGLE
                p_formato.space_after = 0
                p_formato.space_before = 0

            #
            # ---------------------------------------------- Adiciona Proposição
            #
            flag_veto_parcial = dict_projeto[item][6]
            flag_veto_total = dict_projeto[item][7]
            flag_redacao_vencido = dict_projeto[item][8]
            flag_redacao_final = dict_projeto[item][9]
            p = documento_web.add_paragraph()
            p_formato = p.paragraph_format
            p_formato.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p_formato.line_spacing = 1.0
            p_formato.line_spacing_rule = WD_LINE_SPACING.SINGLE
            p_formato.space_after = 0
            p_formato.space_before = 0
            #
            # --------------------------------------------adcionando o ordinal da proposição
            run = p.add_run(texto_ordinal_projeto)
            run.font.color.rgb = RGBColor(dict_cores['CINZA'][0], dict_cores['CINZA'][1], dict_cores['CINZA'][2])
            #
            if flag_veto_parcial == 'Sim' or flag_veto_total == 'Sim':
                lista_aux = dict_chave_projeto[item]
                #
                # --------------------------------------------Link do veto (Parcial ou Total)
                dict_aux = lista_aux[0]
                cor_chave = dict_aux[0][2]
                link_www3 = dict_aux[0][5]
                link_alerjln1 = dict_aux[0][6]
                texto_chave = dict_aux[0][0]
                run = p.add_run(texto_chave)
                run.font.color.rgb = RGBColor(dict_cores[cor_chave][0], dict_cores[cor_chave][1],
                                              dict_cores[cor_chave][2])
                run.bold = True
                # FALTA INCLUIR HIPERLINK
                # --------------------------------------------Link do Autógrafo
                dict_aux = lista_aux[1]
                link_www3 = dict_aux[0][5]
                link_alerjln1 = dict_aux[0][6]
                # ---------------------------------------------Adicionando identificação do projeto
                run = p.add_run(' APOSTO AO ')
                pos_inicial = str.find(texto_projeto, 'AO') + 3
                texto_projeto = texto_projeto[pos_inicial:]
                run = p.add_run(texto_projeto)
                pos_final = str.find(texto_projeto, 'Nº') - 1
                tipo_projeto = texto_projeto[0:pos_final].strip()
                texto_tipo_projeto = dict_tipo_projeto.get(tipo_projeto, 'Não')
                if texto_tipo_projeto != 'Não':
                    cor_projeto = dict_tipo_projeto[tipo_projeto][1]
                else:
                    cor_projeto = 'VERMELHA'
                run.font.color.rgb = RGBColor(dict_cores[cor_projeto][0], dict_cores[cor_projeto][1],
                                              dict_cores[cor_projeto][2])
                run.bold = True
                # FALTA INCLUIR HIPERLINK do veto (texto do ofício) e da proposição (texto do outorgado)
                # ---------------------------------------------Adicionando icone da "cartinha"
                nome_icone = 'img' + item[4:6] + '.png'
                run = p.add_run(' ')
                run.add_picture(nome_icone)
                # FALTA INCLUIR HIPERLINK cartinha (projeto raiz)
            elif flag_redacao_vencido == 'Sim' or flag_redacao_final == 'Sim':
                run = p.add_run(texto_projeto)
                pos_final = str.find(texto_projeto, 'Nº') - 1
                tipo_projeto = texto_projeto[0:pos_final].strip()
                texto_tipo_projeto = dict_tipo_projeto.get(tipo_projeto, 'Não')
                if texto_tipo_projeto != 'Não':
                    cor_projeto = dict_tipo_projeto[tipo_projeto][1]
                else:
                    cor_projeto = 'VERMELHA'
                run.font.color.rgb = RGBColor(dict_cores[cor_projeto][0], dict_cores[cor_projeto][1],
                                              dict_cores[cor_projeto][2])
                run.bold = True
                # FALTA INCLUIR HIPERLINK discurso do vencido
                # ---------------------------------------------Adicionando icone da "cartinha"
                nome_icone = 'img' + item[4:6] + '.png'
                run = p.add_run(' ')
                run.add_picture(nome_icone)
                # FALTA INCLUIR HIPERLINK cartinha (projeto raiz)
            else:
                run = p.add_run(texto_projeto)
                pos_final = str.find(texto_projeto, 'Nº') - 1
                tipo_projeto = texto_projeto[0:pos_final].strip()
                texto_tipo_projeto = dict_tipo_projeto.get(tipo_projeto, 'Não')
                if texto_tipo_projeto != 'Não':
                    cor_projeto = dict_tipo_projeto[tipo_projeto][1]
                else:
                    cor_projeto = 'VERMELHA'
                run.font.color.rgb = RGBColor(dict_cores[cor_projeto][0], dict_cores[cor_projeto][1],
                                              dict_cores[cor_projeto][2])
                run.bold = True

            # ---------------------------------------------Adicionando o termo de autoria
            run = p.add_run(', de autoria ')
            #
            # ---------------------------------------------- Lista os autores da proposição
            #
            possui_autores = dict_autores.get(item, 'Não')
            if possui_autores != 'Não':
                sexo_calculo = 0
                texto_autores = ''
                cont_autores = 0
                total_autores = len(dict_autores[item])
                for deputado in dict_autores[item]:
                    sexo_calculo = sexo_calculo + deputado[2]
                    cont_autores = cont_autores + 1
                    texto_autores = texto_autores + ' ' + deputado[0]  # deputado[1] (Significa "Sim" ou "Não" para achar o nome) - Colocar com risco
                    if total_autores > 1:
                        if cont_autores == total_autores - 1:
                            texto_autores = texto_autores + ' e'
                        elif cont_autores < total_autores:
                            texto_autores = texto_autores + ','
                if sexo_calculo == 0:
                    if total_autores > 1:
                        texto_autores = 'das Deputadas' + texto_autores
                    else:
                        texto_autores = 'da Deputada' + texto_autores
                else:
                    if total_autores > 1:
                        texto_autores = 'dos Deputados' + texto_autores
                    else:
                        texto_autores = 'do Deputado' + texto_autores
                texto_autores = texto_autores.strip()
            #
            # ---------------------------------------------Adicionando autores
            run = p.add_run(texto_autores)

            #
            # ---------------------------------------------Adicionando a ementa
            run = p.add_run(', que')
            valor_ementa = dict_chave_ementa.get(item, 'Não')
            if valor_ementa != 'Não':
                dict_ementa_aux = dict_chave_ementa[item]
                pos_inicial = 0
                pos_final = len(texto_ementa)
                for id in dict_ementa_aux:
                    pos_chave = id
                    texto_chave = dict_ementa_aux[id][0]
                    cor_chave = dict_ementa_aux[id][2]
                    flag_tem_link = dict_ementa_aux[id][3]
                    link_www3 = dict_ementa_aux[id][5]
                    link_alerjln1 = dict_ementa_aux[id][6]
                    run = p.add_run(texto_ementa[pos_inicial:pos_chave])
                    run.font.size = Pt(10)
                    run.bold = True
                    run = p.add_run(texto_chave[0:len(texto_chave)])
                    run.font.color.rgb = RGBColor(dict_cores[cor_chave][0], dict_cores[cor_chave][1],
                                                  dict_cores[cor_chave][2])
                    run.bold = True
                    run.font.size = Pt(10)
                    pos_inicial = pos_chave + len(texto_chave)
                run = p.add_run(texto_ementa[pos_inicial:pos_final])
                run.font.size = Pt(10)
                run.bold = True
            else:
                run = p.add_run(texto_ementa)
                run.bold = True
            #
            # ---------------------------------------------- Lista os pareceres da proposição
            #

            texto_parecer = dict_parecer.get(item, 'Não')
            if texto_parecer != 'Não':
               p = documento_web.add_paragraph()
               texto_parecer = texto_parecer.title()
               dict_parecer_aux = dict_chave_parecer[item]
               p = documento_web.add_paragraph()
               p_formato = p.paragraph_format
               p_formato.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
               p_formato.line_spacing = 1.0
               p_formato.line_spacing_rule = WD_LINE_SPACING.SINGLE
               p_formato.space_after = 0
               p_formato.space_before = 0
               if len(dict_parecer_aux)>1:
                  run = p.add_run('PARECERES: ')
               else:
                  run = p.add_run('PARECER: ')
               run.bold = True
               run.font.size = Pt(10)
               run = p.add_run('da Comissão ')
               run.font.size = Pt(10)
               pos_inicial = 0
               pos_final = len(texto_parecer)
               for id in dict_parecer_aux:
                   pos_chave = id
                   texto_chave = dict_parecer_aux[id][0]
                   cor_chave = dict_parecer_aux[id][2]
                   flag_tem_link = dict_parecer_aux[id][3]
                   link_www3 = dict_parecer_aux[id][5]
                   link_alerjln1 = dict_parecer_aux[id][6]
                   run = p.add_run(texto_parecer[pos_inicial:pos_chave])
                   run.font.size = Pt(10)
                   run = p.add_run(texto_chave[0:1])
                   run.font.size = Pt(10)
                   tamanho = len(texto_chave)
                   ponto_virgula = 1
                   if texto_chave[tamanho-1:tamanho] ==';' or texto_chave[tamanho-1:tamanho] =='.':
                      run = p.add_run(texto_chave[1:len(texto_chave)-1])
                   else:
                      tamanho = len(texto_chave)
                      run = p.add_run(texto_chave[1:len(texto_chave)])
                      ponto_virgula = 0
                   run.font.color.rgb = RGBColor(dict_cores[cor_chave][0], dict_cores[cor_chave][1],
                                                 dict_cores[cor_chave][2])
                   run.bold = True
                   run.font.size = Pt(10)
                   pos_inicial = pos_chave+len(texto_chave)-ponto_virgula
               run = p.add_run(texto_parecer[pos_inicial:pos_final])
               run.font.size = Pt(10)

            #
            # ---------------------------------------------- Lista os novos pareceres da proposição
            #

            texto_novo_parecer = dict_novo_parecer.get(item, 'Não')
            if texto_novo_parecer != 'Não':
                p = documento_web.add_paragraph()
                texto_novo_parecer = texto_novo_parecer.title()
                dict_parecer_aux = dict_chave_novo_parecer[item]
                p = documento_web.add_paragraph()
                p_formato = p.paragraph_format
                p_formato.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p_formato.line_spacing = 1.0
                p_formato.line_spacing_rule = WD_LINE_SPACING.SINGLE
                p_formato.space_after = 0
                p_formato.space_before = 0
                if len(dict_parecer_aux) > 1:
                    run = p.add_run('NOVOS PARECERES: ')
                else:
                    run = p.add_run('NOVO PARECER: ')
                run.bold = True
                run.font.size = Pt(10)
                expressao_regular = '[À|A][S]?[" "]{1,}EMENDA[S]?'
                prog = re.compile(expressao_regular, re.IGNORECASE)
                flag_tem_emenda = 'Não'
                achou = prog.search(texto_novo_parecer)
                if achou:
                   flag_tem_emenda = 'Sim'
                run.font.size = Pt(10)
                pos_inicial = 0
                pos_final = len(texto_novo_parecer)
                cont_parecer = 1
                for id in dict_parecer_aux:
                    pos_chave = id
                    texto_chave = dict_parecer_aux[id][0]
                    cor_chave = dict_parecer_aux[id][2]
                    flag_tem_link = dict_parecer_aux[id][3]
                    link_www3 = dict_parecer_aux[id][5]
                    link_alerjln1 = dict_parecer_aux[id][6]
                    run = p.add_run(texto_novo_parecer[pos_inicial:pos_chave])
                    run.font.size = Pt(10)
                    tamanho = len(texto_chave)
                    ponto_virgula = 1
                    if texto_chave[tamanho-1:tamanho] == ';' or texto_chave[tamanho-1:tamanho] == '.':
                        run = p.add_run(texto_chave[0:1])
                        run.font.size = Pt(10)
                        run = p.add_run(texto_chave[1:len(texto_chave) - 1])
                        run.font.size = Pt(10)
                    else:
                        if flag_tem_emenda == 'Sim':
                           run = p.add_run(texto_chave[0:len(texto_chave)])
                           run.font.size = Pt(10)
                        else:
                            run = p.add_run(texto_chave[1:len(texto_chave)])
                            run.font.size = Pt(10)
                        ponto_virgula = 0
                    run.font.color.rgb = RGBColor(dict_cores[cor_chave][0], dict_cores[cor_chave][1],
                                                  dict_cores[cor_chave][2])
                    run.bold = True
                    run.font.size = Pt(10)
                    if cont_parecer == 1 and flag_tem_emenda != 'Sim':
                        run = p.add_run('da Comissão ')
                        run.font.size = Pt(10)
                        pos_chave = pos_chave + 12
                    cont_parecer = cont_parecer + 1
                    run.font.size = Pt(10)
                    pos_inicial = pos_chave + len(texto_chave) - ponto_virgula
                run = p.add_run(texto_novo_parecer[pos_inicial:pos_final])
                run.font.size = Pt(10)

            #
            # ---------------------------------------------- Lista os relatores da proposição
            #
            texto_relator = ''
            possui_relator = dict_relator.get(item, 'Não')
            if possui_relator != 'Não':
                p = documento_web.add_paragraph()
                p = documento_web.add_paragraph()
                p_formato = p.paragraph_format
                p_formato.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p_formato.line_spacing = 1.0
                p_formato.line_spacing_rule = WD_LINE_SPACING.SINGLE
                p_formato.space_after = 0
                p_formato.space_before = 0
                sexo_calculo = 0
                cont_relator = 0
                total_relator = len(dict_relator[item])
                for deputado in dict_relator[item]:
                    sexo_calculo = sexo_calculo + deputado[2]
                    cont_relator = cont_relator + 1
                    texto_relator = texto_relator + ' ' + deputado[0]  # deputado[1] (Significa "Sim" ou "Não" para achar o nome) - Colocar com risco
                    if total_relator > 1:
                        if cont_relator == total_relator - 1:
                            texto_relator = texto_relator + ' e'
                        elif cont_relator < total_relator:
                            texto_relator = texto_relator + ','
                if sexo_calculo == 0:
                    if total_relator > 1:
                        run = p.add_run('RELATORAS: ')
                        run.font.size = Pt(10)
                        run.bold = True
                        run = p.add_run('Deputadas')
                        run.font.size = Pt(10)
                    else:
                        run = p.add_run('RELATORA: ')
                        run.font.size = Pt(10)
                        run.bold = True
                        run = p.add_run('Deputada')
                        run.font.size = Pt(10)
                else:
                    if total_relator > 1:
                        run = p.add_run('RELATORES: ')
                        run.font.size = Pt(10)
                        run.bold = True
                        run = p.add_run('Deputados')
                        run.font.size = Pt(10)
                    else:
                        run = p.add_run('RELATOR: ')
                        run.font.size = Pt(10)
                        run.bold = True
                        run = p.add_run('Deputado')
                        run.font.size = Pt(10)
                texto_relator = texto_relator + '.'
                #
                # ---------------------------------------------Adicionando relatores
                run = p.add_run(texto_relator)
                run.font.size = Pt(10)
                #
            # ---------------------------------------------- Lista as comissões com pendentes de parecer
            #
            valor_pendente_parecer = dict_comissao_link.get(item, 'Não')
            if valor_pendente_parecer != 'Não':
                p = documento_web.add_paragraph()
                p = documento_web.add_paragraph()
                p_formato = p.paragraph_format
                p_formato.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p_formato.line_spacing = 1.0
                p_formato.line_spacing_rule = WD_LINE_SPACING.SINGLE
                p_formato.space_after = 0
                p_formato.space_before = 0
                run = p.add_run('(PENDENDO DE PARECER: ')
                run.bold = True
                run.font.size = Pt(10)
                run = p.add_run('da Comissão ')
                run.font.size = Pt(10)
                lista_aux = dict_comissao_link[item]
                cont_comissao = 0
                total_comissao = len(lista_aux)
                texto_comissao = ''
                for id in lista_aux:
                    cont_comissao = cont_comissao + 1
                    nome_comissao = id[0]
                    link_www3 = id[1]
                    link_alerjln1 = id[2]
                    if total_comissao > 1:
                        if cont_comissao == total_comissao - 1:
                            run = p.add_run(nome_comissao)
                            run.font.size = Pt(10)
                            run = p.add_run(' e ')
                            run.font.size = Pt(10)
                        elif cont_comissao < total_comissao:
                            run = p.add_run(nome_comissao)
                            run.font.size = Pt(10)
                            run = p.add_run('; ')
                            run.font.size = Pt(10)
                        else:
                            run = p.add_run(nome_comissao)
                            run.font.size = Pt(10)
                    else:
                        run = p.add_run(nome_comissao)
                        run.font.size = Pt(10)

                valor_chave_pendente_parecer = dict_chave_parecer_pendente.get(item, 'Não')
                if valor_chave_pendente_parecer != 'Não':
                    for id in dict_chave_parecer_pendente[item]:
                        lista_aux = dict_chave_parecer_pendente[item][id]
                        texto_chave = lista_aux[0]
                        cor_chave = lista_aux[2]
                        flag_tem_link = lista_aux[3]
                        link_www3 = lista_aux[5]
                        link_alerjln1 = lista_aux[6]
                        run = p.add_run(',')
                        run.font.size = Pt(10)
                        if texto_chave == 'EMENDAS DE PLENÁRIO':
                            run = p.add_run(' às ')
                            run.font.size = Pt(10)
                            run = p.add_run(texto_chave)
                            run.font.size = Pt(10)
                            run.font.color.rgb = RGBColor(dict_cores[cor_chave][0], dict_cores[cor_chave][1],
                                                          dict_cores[cor_chave][2])
                            run.bold = True
                            run = p.add_run('.)')
                            run.font.size = Pt(10)
                        elif texto_chave == 'EMENDA DE PLENÁRIO':
                            run = p.add_run(' à ')
                            run.font.size = Pt(10)
                            run = p.add_run(texto_chave)
                            run.font.size = Pt(10)
                            run.font.color.rgb = RGBColor(dict_cores[cor_chave][0], dict_cores[cor_chave][1],
                                                          dict_cores[cor_chave][2])
                            run.bold = True
                            run = p.add_run('.)')
                            run.font.size = Pt(10)
                else:
                    run = p.add_run('.)')
                    run.font.size = Pt(10)
        #print(lista_atributos_sessao)
	#------definir os nomes dos aquivos de saida
        numero_sessao = lista_atributos_sessao[0].split(" ")[0][:2]
        tipo_sessao = lista_atributos_sessao[0] #.split(" ")[2]
        ano = lista_atributos_sessao[1].split(" ")[4]
        mes = str(TMesAno.obterNumero(lista_atributos_sessao[1].split(" ")[2])).zfill(2)
        dia, hora = lista_atributos_sessao[1].split(" ")[0] , lista_atributos_sessao[3].split("h")[0] +lista_atributos_sessao[3].split("h")[1]
        if(  re.search('EXTRA',tipo_sessao,re.IGNORECASE)):
            tipo_sessao = "SE"
        else:
            tipo_sessao = "SO"
        #print(lista_atributos_sessao)
        #print(numero_sessao + " " + tipo_sessao + " " + ano + " " + mes + " " + dia_hora ) 
        
        #nome_arq_sessao = lista_atributos_sessao[0]+'-'+lista_atributos_sessao[1]+'-Notes'+'.docx'
        nome_arq_sessao = "SESSAO_" + id_processamento + "_" + numero_sessao + "_" +  tipo_sessao + "_" + ano + mes +  dia + "_" + hora + "_Notes.docx"
        lista_arquivos.append(nome_arq_sessao)
        lista_documentos[cont_sessao-1].save("upload/" + nome_arq_sessao)
    #nome_doc_pauta = "PAUTAS" + '-' + lista_atributos_sessao[1] + '-Web'+'.docx'
    nome_doc_pauta = "PAUTAS_" + id_processamento + "_" + ano + mes + dia + "_Web.docx"
    #nome_doc_pauta = nome_doc_pauta.replace(" ", "_")
    documento_web.save("upload/" + nome_doc_pauta)
    lista_arquivos.append(nome_doc_pauta)
    return lista_arquivos

def busca_link_lei_federal(texto):
    #Avaliando de forma mais rigorosa o texto para extrair número da lei, dia, mês e ano de promulgação da lei
    url_raiz = 'https://www.lexml.gov.br/urn/urn:lex:br:federal:lei:'

    lista_meses_ano = ['janeiro',
                       'fevereiro',
                       'março',
                       'abril',
                       'maio',
                       'junho',
                       'julho',
                       'agosto',
                       'setembro',
                       'outubro',
                       'novembro',
                       'dezembro']
    expressao_regular = 'LEI[" "]{1,}FEDERAL.*[0-9]{1,3}[.]?[0-9]{1,3}?[0-9]{1,3}.*[DE][" "]{1,}[1-2][.]?[0-9]{3,3}'
    prog = re.compile(expressao_regular, re.IGNORECASE)
    achou = prog.search(texto)
    ano_lei='0000'
    mes_lei='00'
    dia_lei='00'
    numero_lei='0'
    link_www3 = link_alerjln1 =''
    if achou:
        #-------------------------------------------Identificando o ano
        pos_inicial = achou.span(0)[0]
        pos_final = achou.span(0)[1]
        ano_lei = texto[pos_final-4:pos_final]
        ano_lei = re.sub('[^0-9]', '', ano_lei)
        #-------------------------------------------Identificando o número do projeto
        pos_inicial = texto.find('FEDERAL') + 8
        pos_final = texto.find('DE',pos_inicial)
        numero_lei = texto[pos_inicial:pos_final]
        numero_lei = re.sub('[^0-9]', '', numero_lei)
        # --------------------------------------- --Identificando o dia do projeto
        pos_inicial = pos_final+2
        pos_final = texto.find('DE', pos_inicial)
        if pos_final!=-1:
           dia_lei = texto[pos_inicial:pos_final]
           dia_lei = re.sub('[^0-9]', '', dia_lei)
        #-------------------------------------------Identificando o mês do projeto
        cont_mes=0
        for item in lista_meses_ano:
            cont_mes=cont_mes+1
            if texto.find(item.upper())!=-1:
               mes_lei = str(cont_mes)
               if len(mes_lei) == 1:
                   mes_lei = '0'+mes_lei
                   break
        link_alerjln1=url_raiz+ano_lei+'-'+mes_lei+'-'+dia_lei+';'+numero_lei
        link_www3 =link_alerjln1

    return(link_alerjln1,link_www3)

def busca_link_lei_estadual(texto):
    expressao_regular_1 = 'LEI[" "]{1,}.*[0-9]{1,3}[.]?[0-9]{1,3}?[0-9]{1,3}.*/{1,}[1-2][.]?[0-9]{3,3}'
    expressao_regular_2 = 'LEI[" "]{1,}[^FEDERAL].*[0-9]{1,3}[.]?[0-9]{1,3}?[0-9]{1,3}.*[DE][" "]{1,}[1-2][.]?[0-9]{3,3}'
    #url_raiz = 'http://191.34.56.77:5101/lei/aaaa;nnnnn
    url_raiz = 'https://lotusnotesapp.herokuapp.com/lei/'
    url_busca=''
    ano_lei = '0000'
    numero_lei = '0'
    link_www3 = link_alerjln1 = ''
    prog = re.compile(expressao_regular_1, re.IGNORECASE)
    achou = prog.search(texto)
    if achou:
        # -------------------------------------------Identificando o ano
        pos_inicial = achou.span(0)[0]
        pos_final = achou.span(0)[1]
        ano_lei = texto[pos_final - 4:pos_final]
        ano_lei = re.sub('[^0-9]', '', ano_lei)
        # -------------------------------------------Identificando o número do projeto
        pos_final = texto.find('/', pos_inicial)
        numero_lei = texto[pos_inicial:pos_final]
        numero_lei = re.sub('[^0-9]', '', numero_lei)
        # --------------------------------------- --Identificando o dia do projeto
        #Pesquisar serviço do Pavão para identificar o link
        url_busca = url_raiz + ano_lei +";"+numero_lei
    else:
        prog = re.compile(expressao_regular_2, re.IGNORECASE)
        achou = prog.search(texto)
        if achou:
            # -------------------------------------------Identificando o ano
            pos_inicial = achou.span(0)[0]
            pos_final = achou.span(0)[1]
            ano_lei = texto[pos_final - 4:pos_final]
            ano_lei = re.sub('[^0-9]', '', ano_lei)
            # -------------------------------------------Identificando o número do projeto
            pos_inicial = texto.find('Nº')+2
            if pos_inicial != -1:
                pos_inicial = texto.find('Nº') + 2
            else:
                pos_inicial = texto.find('LEI') + 3
            pos_final = texto.find('DE')-1
            numero_lei = texto[pos_inicial:pos_final]
            numero_lei = re.sub('[^0-9]', '', numero_lei)
    # Pesquisar serviço do Pavão para identificar o link
    texto_url=''
    url_busca = url_raiz + ano_lei + ";" + numero_lei
    page = requests.get(url_busca)
    if page.status_code == 200:
       soup = BeautifulSoup(page.text, 'html.parser')
       texto_url = soup.prettify()
    link_alerj, link_www3 = busca_link_raiz(texto_url)
    return (link_alerj, link_www3)


def busca_link_raiz(texto):
    link_alerj=''
    link_www3=''
    expressao_regular="link notes:"
    prog = re.compile(expressao_regular,re.IGNORECASE)
    achou = prog.search(texto)
    if achou:
        pos_inicial = achou.span(0)[0]
        pos_inicial_link_alerj = texto.find("http", pos_inicial)
        pos_final_link_alerj = texto.find("<br/>", pos_inicial_link_alerj)
        link_alerj = texto[pos_inicial_link_alerj:pos_final_link_alerj]
        link_alerj=link_alerj.replace('\n', '')
        link_alerj = link_alerj.strip()

    expressao_regular="link WWW3:"
    prog = re.compile(expressao_regular,re.IGNORECASE)
    achou = prog.search(texto)
    if achou:
        pos_inicial = achou.span(0)[0]
        pos_inicial_link_www3 = texto.find("http", pos_final_link_alerj)
        pos_final_link_www3 = texto.find("<br/>", pos_inicial_link_www3)
        link_www3 = texto[pos_inicial_link_www3:pos_final_link_www3]
        link_www3 = link_www3.replace('amp;', '')
        link_www3 = link_www3.replace('\n', '')
        link_www3 = link_www3.strip()
    return(link_alerj,link_www3)

def busca_link_parecer(expressao_1,expressao_2,texto):
    pos_ajuste = expressao_1.find("COMISSÃO")
    if pos_ajuste!= -1:
       expressao_1=expressao_1[pos_ajuste+8:]
    expressao_1 = expressao_1.replace(',', '[,]?')
    expressao_regular=expressao_1+'.*'+expressao_2
    prog = re.compile(expressao_regular,re.IGNORECASE)
    link_alerj=''
    link_www3=''
    achou = prog.search(texto)
    if achou:
        pos_inicial = achou.span(0)[0]
        pos_inicial_link_alerj = texto.find("http", pos_inicial)
        pos_final_link_alerj = texto.find("<br/>", pos_inicial_link_alerj)
        pos_inicial_link_www3 = texto.find("http", pos_final_link_alerj)
        pos_final_link_www3 = texto.find("<br/>", pos_inicial_link_www3)
        link_alerj = texto[pos_inicial_link_alerj:pos_final_link_alerj]
        link_www3 = texto[pos_inicial_link_www3:pos_final_link_www3]
        link_www3 = link_www3.replace('amp;', '')
        link_alerj=link_alerj.replace('\n', '')
        link_alerj = link_alerj.strip()
        link_www3 = link_www3.replace('\n', '')
        link_www3 = link_www3.strip()
    return(link_alerj,link_www3)

def busca_deputado(texto):
    lista = []
    lista_tratada = []
    lista_deputado = []
    texto_achou = ''
    nome_deputado=''
    achou = re.search("^.*D[A|O]S? DEPUTAD[A|O]S?|^.*DEPUTAD[A|O]S?", texto)
    if achou:
        texto_achou = achou.group(0)
        texto_aux=texto[len(texto_achou):]
        texto_aux=re.sub(r'( E )', r', ', texto_aux)
        lista=(texto_aux.split(','))
    for item in lista:
        nome_deputado=item.strip()
        pos = nome_deputado.find('.',len(nome_deputado)-1)
        if pos != -1:
            nome_deputado=nome_deputado[0:len(nome_deputado)-1]
        achou, sexo_calculo = verifica_nome_sexo_deputado(nome_deputado)
        lista_deputado = []
        nome_deputado = nome_deputado.title()
        nome_deputado = re.sub(r'( Do )', r' do ', nome_deputado)
        nome_deputado = re.sub(r'( Da )', r' da ', nome_deputado)
        nome_deputado = re.sub(r'( De )', r' de ', nome_deputado)
        nome_deputado = re.sub(r'( Dos )', r' dos ', nome_deputado)
        nome_deputado = re.sub(r'( Das )', r' das ', nome_deputado)
        lista_deputado.append(nome_deputado)
        lista_deputado.append(achou)
        lista_deputado.append(sexo_calculo)
        lista_tratada.append(lista_deputado)
    return(lista_tratada)

def busca_link_comissao_pendente(texto_comissao_pendente,dict_comissao_dominio):
    lista_comissao_aux =[]
    lista_comissão_pendente = []
    texto_comissao_pendente=texto_comissao_pendente.replace(")", "")
    texto_aux =''
    achou_comissao =''
    pos_inicial = 0
    if re.search('^.*[À|A][S}?.*EMENDA[S]? DE PLEN[A|Á]RIO', texto_comissao_pendente): #Retira a referência de parecer pendente à(s) emenda(s) de planário
       pos_final = texto_comissao_pendente.find(", À")+1
    else:
        pos_final = len(texto_comissao_pendente)
    lista_comissao_aux = texto_comissao_pendente[pos_inicial:pos_final].split(';')
    for item in lista_comissao_aux:
        lista_atributos = []
        if item[0:2] == ' E':
            texto_aux = item[2:len(item)+1]
            texto_aux = texto_aux.strip()
        else:
            texto_aux = item.strip()
        chave = texto_aux.upper()
        chave = remove_acentos(chave)
        chave = re.sub(',', '', chave)
        texto_parecer = ''
        achou_comissao = dict_comissao_dominio.get(chave, 'Não')
        if achou_comissao != 'Não':
           lista_atributos = dict_comissao_dominio[chave]
        else:
           lista_atributos.append(texto_aux)
           lista_atributos.append('')
           lista_atributos.append('')
        lista_comissão_pendente.append(lista_atributos)
    return(lista_comissão_pendente)

def busca_palavras_chaves(texto, dict_palavras_chaves,id_projeto,texto_url,dict_comissao_dominio):
   dict_posicao_chave = {}
   dict_posicao_chave_ordenada = {}
   dict_posicao_cor_link = {}
   lista_chave=[]
   lista_ordenada=[]
   lista_achado=[]
   lista_achado_ordenada=[]
   expressao_regular = '' #Padrão de busca da palavra_chave na pauta
   texto_cor = '' #Cor do texto
   flag_tem_link ='Não' #Indica se o parecer precisa ser encontrado na página de tramitação da matéria
   expressao_regular_tramitacao ='' #Padrão de busca da palavra_chave na tramitacao
   pos=0
   for key in dict_palavras_chaves:
       lista_palavras_chaves_aux = dict_palavras_chaves[key]
       palavra_chave = key
       expressao_regular = lista_palavras_chaves_aux[0]
       texto_cor = lista_palavras_chaves_aux[1]
       flag_tem_link = lista_palavras_chaves_aux[2]
       expressao_regular_tramitacao = lista_palavras_chaves_aux[3]
       lista_achado = re.findall(expressao_regular, texto)
       lista_achado_ordenada=sorted(lista_achado)
       item_anterior = ''
       pos_anterior = 0
       pos_final = len(texto)
       for item in lista_achado_ordenada:
           if item != item_anterior:
              pos_inicial = 0
              item_anterior = item
           else:
              pos_inicial = pos +len(item_anterior)
           pos = texto.find(item,pos_inicial,pos_final)
           lista_cor_link =[]
           lista_cor_link.append(texto_cor)
           lista_cor_link.append(flag_tem_link)
           lista_cor_link.append(expressao_regular_tramitacao)
           dict_posicao_cor_link.update({pos:lista_cor_link})
           dict_posicao_chave.update({pos:item})
   for key in dict_posicao_chave:
       lista_chave.append(key)
   lista_ordenada = sorted(lista_chave)
   contador_item = 0
   flag_emenda_plenario = "Não"
   nome_comissao = ''
   pos_inicial = 0
   for item in lista_ordenada:
       lista_busca = []
       contador_item = contador_item + 1
       pos_final = item
       if contador_item == 1:
           if re.search('EMENDA[S]? DE PLEN[A|Á]RIO',dict_posicao_chave[item]):
              nome_comissao = ''
              lista_busca.append(dict_posicao_chave[item])
              lista_busca.append(nome_comissao)
              lista_busca.append(dict_posicao_cor_link[item][0])
              lista_busca.append(dict_posicao_cor_link[item][1])
              lista_busca.append(dict_posicao_cor_link[item][2])
           else:
              nome_comissao = texto[pos_inicial:pos_final]
              pos_inicial = nome_comissao.find("COMISSÃO")
              if pos_inicial!=-1:
                 nome_comissao=nome_comissao[pos_inicial+8:]
              if nome_comissao[0:2] == ' E':
                 nome_comissao = nome_comissao[2:len(nome_comissao)]
              nome_comissao = nome_comissao.strip()
              lista_achado = busca_link_comissao_pendente(nome_comissao, dict_comissao_dominio)
              nome_comissao = lista_achado[0][0]
              lista_busca.append(dict_posicao_chave[item])
              lista_busca.append(nome_comissao)
              lista_busca.append(dict_posicao_cor_link[item][0])
              lista_busca.append(dict_posicao_cor_link[item][1])
              lista_busca.append(dict_posicao_cor_link[item][2])
       else:
           nome_comissao = texto[pos_inicial:pos_final]
           pos_inicial = nome_comissao.find("COMISSÃO")
           if pos_inicial != -1:
               nome_comissao = nome_comissao[pos_inicial + 8:]
           if nome_comissao[0:2] == ' E':
              nome_comissao = nome_comissao[2:len(nome_comissao)]
           nome_comissao = nome_comissao.strip()
           lista_achado = busca_link_comissao_pendente(nome_comissao, dict_comissao_dominio)
           nome_comissao = lista_achado[0][0]
           lista_busca.append(dict_posicao_chave[item])
           lista_busca.append(nome_comissao)
           lista_busca.append(dict_posicao_cor_link[item][0])
           lista_busca.append(dict_posicao_cor_link[item][1])
           lista_busca.append(dict_posicao_cor_link[item][2])
       #------------------------------------------------------------------- Colocar a Busca dos links

       link_www3=''
       link_alerjln1=''
       if lista_busca[3] =='Sim': #Precisa de hiperlink
          if re.search('^LEI FEDERAL.*', dict_posicao_chave[item]): #Constrói link de lei federal
             link_alerjln1,link_www3 = busca_link_lei_federal(dict_posicao_chave[item])
          elif re.search('^LEI.*', dict_posicao_chave[item]): #Constrói link de lei estadual
              link_alerjln1,link_www3 = busca_link_lei_estadual(dict_posicao_chave[item])
          else:
              link_alerjln1,link_www3 = busca_link_parecer(lista_busca[1],lista_busca[4],texto_url)
       lista_busca.append(link_www3)
       lista_busca.append(link_alerjln1)
       dict_posicao_chave_ordenada.update({item: lista_busca})
       pos_inicial=pos_final+len(dict_posicao_chave[item])
   return(dict_posicao_chave_ordenada)

def dia_semana (ano,mes, dia):
   # Retorna o dia da semana

   lista_dias_semana = ['Segunda-feira', 'Terça-feira','Quarta-feira','Quinta-Feira','Sexta-feira','Sábado','Domingo']
   data = date(year=ano, month=mes, day=dia)
   indice_da_semana = data.weekday()
   dia_da_semana = lista_dias_semana[indice_da_semana]
   return(dia_da_semana)

def main(arquivo_entrada,id_processamento):
#
#-------------------------------Declaração das Variáveis-------------------
   dict_sessao = {} # Relação das Sessões contendo: id_sessao, lista_atributo_sessao_aux
   lista_atributo_sessao_aux = []  # Lista auxiliar de atributos da sessao contendo: texto_sessao, texto_data_sessao, texto_dia_semana_sessao, texto_hora_sessao
   lista_chave_projeto_aux = [] # Lista auxiliar de chaves de projetos a serem relacionadas a uma sessao contendo: id_projeto
   lista_situacao_projeto_aux = []  # Lista auxiliar de situações de projetos a serem relacionadas ao projeto contendo: texto_situação
   lista_situacao_projeto_ant = []  # Lista anterior de situações de projetos a serem relacionadas ao projeto contendo: texto_situação
   dict_projeto_sessao ={} # Relação dos Projetos dentro das Sessões contendo: id_sessao, lista_chave_projeto_aux
   lista_atributos_projeto_aux = [] # Lista auxiliar de atributos da sessao contendo: texto_projeto, texto_ordinal_projeto, texto_inclusao_ri, lista_situacao_projeto, texto_autores, texto_ementa,flag_veto_parcial,flag_veto_total_flag_redacao_vencido,flag_redacao_final
   dict_projeto = {} # Relação dos projetos contendo: id_projeto, lista_atributos_projeto_aux
   dict_parecer = {} # Relação dos pareceres contendo: id_projeto, texto_parecer
   dict_novo_parecer = {} # Relação dos novos pareceres contendo: id_projeto, texto_novo_parecer
   dict_comissao_pendente = {} # Relação das comissões pendentes contendo: id_projeto, texto_comissao_pendente
   dict_comissao_dominio ={} #Relação das comissões existentes no dominio do Notes
   dict_comissao_link ={} #Relação das comissões com pendência de parecer e os respectivos links Notes (www3 e alerjln1)
   dict_relator = {} # Relação dos relatores contendo: id_projeto, texto_relator
   dict_chave_projeto = {} #dicionário de palavras-chaves encontradas no texto novos pareceres de um projeto ou outra matéria
   dict_chave_ementa = {} #dicionário de palavras-chaves encontradas no texto novos pareceres de um projeto ou outra matéria
   dict_chave_novo_parecer = {} #dicionário de palavras-chaves encontradas no texto novos pareceres de um projeto ou outra matéria
   dict_chave_parecer = {} #dicionário de palavras-chaves encontradas no texto de pareceres de um projeto ou outra matéria
   dict_chave_parecer_pendente = {}  #dicionário de palavras-chaves encontradas no texto de pareceres pendentes de um projeto ou outra matéria (EMENDA DE PLENÁRIO)
   lista_chave_achada =[] #Lista de palavras-chaves encontrada no texto
   lista_chave_achada_aux =[] #Lista de palavras-chaves encontrada no texto
   lista_palavras_chaves = [] # Lista de palavras-chaves a serem buscada nos textos das partes PARECERES e NOVOS PARECERES contento: texto_chave, texto_RE, texto_cor, flag_tem_link
   dict_palavras_chaves = {} # Dicionário criado a partir da lista_palavras_chaves
   lista_comissao_pendente_achada = [] #Lista de comissões achadas nos textos da parte PENDENDO DE PARECER(ES) contendo: texto_comissao_pendente, flag_tem_link, texto_link, flag_emenda_plenario
   lista_autores = [] #Lista de autores da matéria contendo: nomes dos deputados autores
   dict_autores = {} # Dicionário criado a partir da lista_autores
   lista_relator = [] #Lista de relatores da matéria contendo: nomes dos deputados relatores
   dict_relator = {} # Dicionário criado a partir da lista_relator
   lista_cores = []  # Lista as cores das palavras-chaves contendo: texto cor e lista da RGB da cor
   dict_cores ={} # Dicionário criado a partir da lista chamada lista_cores
   lista_link_raiz =() #Lista com os links raiz (WWW3 e Alerjln1) da proposição
   dict_link_raiz = {} #Dicionário com os links raiz da proposição


   #---------------------------------------------------------------- Variaveis Auxiliares de Sessao

   texto_sessao = ''             # Nome da sessao corrente
   texto_sessao_anterior =''     # Nome da sessao anterior
   texto_data_sessao =''         # Texto com data completa da sessao (dia de mes de ano) 
   texto_dia_semana_sessao =''   # Texto com data completa da sessao (segunda-feira, terça-feira, ...)
   texto_hora_sessao =''         # Texto com hora completa da sessao 
   flag_primeira_sessao = 'Sim'  # Indica se é a primeira sessao lida
   flag_tem_sessao = 'Não'       # Indica se foi achado uma sessao
   flag_tem_regime = 'Não'       # Indica se foi achado tipo de regime de uma sessão
   flag_tem_discussao = 'Não'    # Indica se foi achado tipo de discussão de uma sessão
   cont_sessao=0                 # Quantidade de sessões lidas
   numero_mes_sessao = 0         # Número do mês da sessão
   numero_ano_sessao = 0         # Número do ano da sessão
   numero_dia_sessão = 0         # Número do dia da sessão

#---------------------------------------------------------------- Variaveis Auxiliares de Materia (Processo ou outros)

   id_projeto=''              # Texto que descreve o código interno no notes projeto ou outra matéria
   id_projeto_anterior=''     # Texto que descreve o código interno no notes projeto ou outra matéria anterior
   seq_numero_projeto=''      # Sequencia ordinal do projeto ou da outra matéria 
   texto_numero_projeto=''    # Texto que descreve o número do projeto ou outra matéria (rótulo completo sobre o qual será incluido hiperlink)
   texto_ano_projeto=''       # Texto que descreve o ano do projeto ou outra matéria
   texto_ementa=''            # Texto que descreve o projeto ou outra matéria
   numero_ordinal_projeto=0   # Ordem em que a materia aparece dentro da sessão
   texto_ordinal_projeto=''   # Texto do ordinal da matéria na sessão
   texto_inclusao_ri=''       # Texto de justificativa de inclusão de materia baseado em Regimento Interno
   texto_situacao = ''        # Texto que descreve a situação em que se encontra o projeto (pex. "Em 2a Votação", "Em discussão"..)
   texto_autores = ''         # Texto da autores da matéria        
   texto_parecer =''          # Texto do(s) parecer(s) comissões da matéria
   texto_novo_parecer =''     # Texto do(s) de novo(s) parecer(s) comissões da matéria
   texto_relator = ''         # Texto do relator da matéria
   texto_tipo_projeto =''     # Texto do tipo da matéria
   texto_comissao_pendente=''    # Texto da comissão pendente da matéria
   flag_tem_autoria = 'Não'   # Indica se foi achado a autoria de uma matéria
   flag_tem_numero = 'Não'    # Indica se foi achado número para a matéria
   flag_tem_parecer ='Não' # Indica se foi achado parecer(es)
   flag_tem_novo_parecer ='Não' # Indica se foi achado novo(s) parecer(es)
   flag_tem_pendente = 'Não'  # Indica se foi achado pendente de comissão
   flag_tem_relator = 'Não'   # Indica se foi achado relator da materia
   flag_passou_processo = "Não" # Indica novo agrupamento de situação de projeto
   cont_projeto=0             # Quantidade de materias lidas totais (Auxilia em produzir uma chave única para matérias sem número)
   cont_projeto_ordem=0       # Quantidade de materias lidas totais (Auxiliar para definir o ordinal da matéria dentro da sessão)
   cont_situacao=0
   flag_primeiro_processo = 'Sim'  # Indica se é o primeiro processo
   flag_ultimo_marcador = 'Não' # Indica se passou pela última palavra reservada
   flag_veto_total = 'Não' # Indica que há veto total da matéria
   flag_veto_parcial = 'Não' # Indica que há veto parcial da matéria
   flag_redacao_vencido = 'Não' # Indica que há redação do vencido
   flag_redacao_final = 'Não' # Indica que há redação final
   flag_tem_serviço_http = 'Não' # Indica que o serviço de internet está disponível
   #url_projeto ='http://mphome.dyndns-at-home.com:5101/notes/projeto/'
   #url_projeto ='http://179.83.247.47:5101/notes/projeto/'
   url_projeto = 'https://lotusnotesapp.herokuapp.com/projeto/'
   url_busca = ''

#---------------------------------------------------------------- Variaveis Auxiliares da logica de programação

   pos_numero=0               # Variavel auxiliar que guarda a posição de um caracter uma cadeia
   pos_barra=0                # Variavel auxiliar que guarda a posição de um caracter uma cadeia
   pos_inicio=0                # Variavel auxiliar que guarda a posição de um caracter uma cadeia
   pos_final=0                # Variavel auxiliar que guarda a posição de um caracter uma cadeia
   pos_data=0                 # Variavel auxiliar que guarda a posição de um caracter uma cadeia
   nome_arq_entrada=''        # Nome do arquivo de entrada (Alyne) (Será passado como parâmetro e essa atribuição irá ser retirada)
   input_path = "/home/mendes/Downloads/" # Caminho da pasta que contem o arquivo de origem (Será passado como parâmetro e essa atribuição irá ser retirada)
   out_path = "/home/mendes/Downloads/" # Caminho da pasta que contem o arquivo de destino (Será passado como parâmetro e essa atribuição irá ser retirada)
   
#-----------------------------Tabelas(dicionários) de Domínio (Origem Notes)
#-----------------------------Tabela de Tipo de Projetos

   lista_tipo_projeto = [('PROPOSTA DE EMENDA CONSTITUCIONAL',('01','MARROM_CLARO','img01.png')),
                         ('PROJETO DE LEI COMPLEMENTAR',('02','VERDE_CLARO','img02.png')),
                         ('PROJETO DE LEI',('03','VERDE_ESCURO','img03.png')),
                         ('PROJETO DE DECRETO LEGISLATIVO',('04','ABACATE','img04.png')),
                         ('PROJETO DE RESOLUÇÃO',('05','AZUL_ESCURO','img05.png')),
                         ('INDICAÇÃO LEGISLATIVA',('06','AZUL_CLARO','img07.png')),
                         ('INDICAÇÃO',('07','PRETA','Sem')),
                         ('MENSAGEM',('08','PRETA','Sem')),
                         ('REQUERIMENTO DE INFORMAÇÕES',('09','PRETA','Sem')),
                         ('REQUERIMENTO',('10','MARROM_ESCURO','img10.png')),
                         ('OFÍCIO',('11','PRETA','Sem')),
                         ('DENÚNCIA POR CRIME DE RESPONSABILIDADE',('12','PRETA','Sem')),
                         ('PROCESSO',('13','PRETA','Sem')),
                         ('MOÇÃO',('14','PRETA','Sem')),
                         ('REQUERIMENTO SEM NÚMERO',('15','PRETA','Sem'))]

   dict_tipo_projeto = dict(lista_tipo_projeto)

   lista_cores = [('AMARELA',(255, 192, 0)),
                  ('ROXA',(139,0,139)),
                  ('ROSA',(255, 51, 204)),
                  ('VERMELHA', (192, 0, 0)),
                  ('AZUL_CLARO',(0, 153, 204)),
                  ('AZUL_ESCURO',(5, 99, 193)),
                  ('MARROM_CLARO',(196, 89, 17)),
                  ('MARROM_ESCURO', (128, 98, 16)),
                  ('ABACATE', (0, 128, 128)),
                  ('VERDE_ESCURO',(0, 130, 80)),
                  ('VERDE_CLARO',(0, 176, 80)),
                  ('CINZA',(128,128,128)),
                  ('PRETA',(0,0,0))]

   dict_cores = dict(lista_cores)

#-----------------------------Carrega domínio de comissões parlamentares existentes no Notes

   dict_comissao_dominio = carrega_dominio_comissoes()

#-----------------------------Tabela de Palavras-Chaves 

   #Lista de palavras-chaves definidas para terem cores distintas e terem ou não hiperlinks para tramitações contendo: texto_chave, texto_RE, texto_cor, flag_tem_link

   lista_palavras_chaves = [('EMENDA(S) DE PLENÁRIO',('EMENDA[S]? DE PLEN[A|Á]RIO?','AMARELA','Sim','OBJETO PARA APRECIAÇÃO.*EMENDA')),
                            ('FAVORÁVEL',(", FAVOR[Á|A]VEL[.|;]","PRETA","Não",'FAVORÁVEL')),
                            ('VETO TOTAL',("^VETO TOTAL","ROXA","Sim",'OFÍCIO ORIGEM.*VETO TOTAL')),
                            ('VETO PARCIAL',("^VETO PARCIAL","ROSA","Sim",'OFÍCIO ORIGEM.*VETO PARCIAL')),
                            ('AUTÓGRAFO', ("^AUTÓGRAFO", "PRETA", "Sim", 'TRAMITAÇÃO DE AUTÓGRAFO.*PODER EXECUTIVO')), #A cor herdade será a do tipo de proposição
                            ('REDAÇÃO DO VENCIDO',('REDAÇÃO DO VENCIDO','PRETA', 'Sim','REDAÇÃO DO VENCIDO.*COMISSÃO DE REDAÇÃO')), # A cor herdade será a do tipo de proposição
                            ('REDAÇÃO FINAL',('REDAÇÃO FINAL','PRETA', 'Sim','REDAÇÃO FINAL.*COMISSÃO DE REDAÇÃO')),# A cor herdade será a do tipo de proposição
                            ('PELA LEGALIDADE, COM EMENDA(S)',(", PELA LEGALIDADE, COM EMENDA[S]?[;|.]","AZUL_CLARO","Sim",'PELA LEGALIDADE[,]? COM EMENDA')),
                            ('PELA CONSTITUCIONALIDADE, COM EMENDA(S)',(", PELA CONSTITUCIONALIDADE[,]? COM EMENDA[S]?[;|.]","AZUL_CLARO", "Sim",'PELA CONSTITUCIONALIDADE[,]? COM EMENDA')),
                            ('PRAZO FINAL',('^PRAZO FINAL: [0-9]{2,2}/[0-9]{2,2}/[0-9]{4,4}','VERMELHA','Não','PRAZO FINAL')),
                            ('CONTRÁRIO',(', CONTR[Á|A]RIO[;|.]','VERMELHA','Sim','CONTRÁRIO')),
                            ('PELA CONSTITUCIONALIDADE',(", PELA CONSTITUCIONALIDADE[,]?[;|.]", 'PRETA', 'Não','PELA CONSTITUCIONALIDADE')),
                            ('LEI FEDERAL', ('LEI FEDERAL.*[1-2].?[0-9][0-9][0-9]', 'AZUL_CLARO', 'Sim','LEI FEDERAL')),
                            ('LEI ESTADUAL', ('LEI Nº [0-9]{1,3}.?[0-9][0-9][0-9]?/[1-2][0-9][0-9][0-9]|LEI[" "]{1,}[^FEDERAL].*[0-9]{1,3}[.]?[0-9]{1,3}?[0-9]{1,3}.*[DE][" "]{1,}[1-2][.]?[0-9]{3,3}', 'AZUL_CLARO', 'Sim','LEI ESTADUAL')),
                            ('FAVORÁVEL, COM A(S) EMENDA(S)',(', FAVORÁVEL,? COM A[S]? EMENDA[S]?','PRETA','Não','FAVORÁVEL COM EMENDA'))]

   dict_palavras_chaves=dict(lista_palavras_chaves)

#-----------------------------Tabelas de Dias da Semana e meses do ano


   lista_meses_ano = ['janeiro',
                      'fevereiro',
                      'março',
                      'abril',
                      'maio',
                      'junho',
                      'julho',
                      'agosto',
                      'setembro',
                      'outubro',
                      'novembro',
                      'dezembro']


#
#-----------------------------Início da Leitura do Documento de Origem (Documento Sem  Formatação)
#
   #arquivo_entrada = 'PAUTA VETOS 24 SETEMBRO.docx' # Deverá ser suprimido (nome do arquivo será recebido por parâmetro)
   #arquivo_entrada = 'Sessões Extra e Ordinária - 18 DE JUNHO_TESTE.docx'
   #arquivo_entrada = 'PAUTA DO DIA 11.docx'
   document = Document(arquivo_entrada) # Carrega o documento de origem (Alyne) a ser lido

   for p in document.paragraphs: #Lendo cada paragrafo Documento de Origem (Documento da Alyne)
      #
      #--------------------------Identificando a Sessão
      #
      achou = re.search("^^.*SESS[AÃ]O.*ORDIN[AÁ]RIA",p.text)
      if achou:
         texto_sessao = achou.group(0)
         #cont_sessao = cont_sessao+1
         achou = re.search("[0-9]{4,4}$",p.text)
         if achou:
            numero_ano_sessao=int(achou.group(0))
         else:
            ano_sessão = ''
         achou = re.search('DIA[" "]{1,}[DE]?.*[0-9][0-9][" "]{1,}[DE].*[1-2][0-9][0-9][0-9]$', p.text)
         if achou:
            pos_numero = achou.span(0)[0] #posição string DIA
            texto_data_sessao = (p.text[pos_numero+3:len(p.text)]).strip()
            if texto_data_sessao[0:2] == 'DE':
                texto_data_sessao=texto_data_sessao[3:len(texto_data_sessao)]
            numero_dia_sessao = int(texto_data_sessao[0:2])
            numero_mes_sessao = 0
            for i in lista_meses_ano:
               numero_mes_sessao = numero_mes_sessao + 1
               mes_ano = i.upper()
               achou = re.search(mes_ano,texto_data_sessao)
               if achou:
                  break
         else:
            texto_data_sessao = ''
            numero_dia_sessao = 0
            numero_mes_sessao = 0
            numero_ano_sessao = 0


         texto_dia_semana_sessao = dia_semana(numero_ano_sessao, numero_mes_sessao, numero_dia_sessao)
         continue
      #
      # --------------------------Identificando a Hora da Sessão
      #
      padrao = "[0-9]{2}[h][0-9]{2}"
      if re.match(padrao, p.text):
         texto_hora_sessao = (p.text).strip()
         continue
      #
      # --------------------------Identificando Matéria incluídas pelo Regimento Interno
      #
      achou = re.search("^.*INCLU[IÍ]DOS.*§",p.text)
      if achou:
         texto_inclusao_ri = p.text.strip()
         continue

      #
      #--------------------------Identificando a situação do Projeto de Lei (colocar num dicionário lendo de arquivo)
      # Colocar chaves em base de dados
      #
      achou=re.search("^.*URG[E|Ê]NCIA|^.*DISCUS{1,2}[A|Ã]O|^.*REDA[Ç|C][A|Ã]O[" "]{1,}FINAL|^.*REDA[Ç|C][A|Ã]O[" "]{1,}DO[" "]{1,}VENCIDO|^.*TRAMITA[Ç|C][A|Ã]O|^.*VOTA[Ç|C][A|Ã]O", p.text)
      if achou:
         cont_situacao=cont_situacao+1
         if flag_passou_processo == 'Sim':
             if cont_situacao == 1:
                lista_situacao_projeto_aux = []
             flag_passou_processo == 'Não'
         lista_situacao_projeto_aux.append(p.text.strip())
         continue

      #
      #--------------------------Identificando a Hora da Sessão
      #
      padrao = "[0-9]{2}[h][0-9]{2}"
      if re.match(padrao, p.text):
         texto_hora_sessao=(p.text).strip()
         #print(id_projeto)
         #print("Achei Hora da Sessão!")
         continue
      #
      #--------------------------Identificando o Projeto ou tipo de matéria
      #
      flag_tem_autoria = 'Sim' if 'AUTORIA' in p.text else 'Não'  # Melhorar a expressão regular
      if flag_tem_autoria =='Sim':
         texto_autores = ''
         texto_ementa = ''
         flag_veto_total = 'Não'  # Indica que há veto total da matéria
         flag_veto_parcial = 'Não'  # Indica que há veto parcial da matéria
         flag_redacao_vencido = 'Não'  # Indica que há redação do vencido
         flag_redacao_final = 'Não'  # Indica que há redação final
         flag_passou_processo = 'Sim'
         cont_situacao =0
         pos_inicio = p.text.find('AUTORIA')+7
         pos_final = p.text.find('QUE',pos_inicio )            #Melhorar a expressão regular pra ficar mais robusto
         texto_autores = p.text[pos_inicio:pos_final-2]
         texto_ementa = p.text[pos_final+3:len(p.text)] #retirando a palavra 'QUE' -1
         cont_projeto=cont_projeto+1
         flag_tem_numero = 'Sim' if 'Nº' in p.text else 'Não' # Melhorar a expressão regular
         if flag_tem_numero == 'Sim':
            pos_numero = str.find(p.text, "Nº")
            pos_barra = str.find(p.text, "/")
            seq_numero_projeto = (p.text[pos_numero + 2:pos_barra]).strip()
            seq_numero_projeto = re.sub('[^0-9]', '', seq_numero_projeto) # Retirar caracter 'A' para projetos com discurso de vencidos
            texto_numero_projeto = p.text[0:pos_barra + 5]
            id_projeto = seq_numero_projeto.zfill(5)  # Ajustando a sequência do número com zeros à esquerda para completar string contendo 5 algarismos
            texto_ano_projeto = p.text[pos_barra + 1:pos_barra + 5]
         else:
            seq_numero_projeto = 'Sem Numero'
            pos_numero = str.find(p.text, ",")
            texto_numero_projeto = p.text[0:pos_numero - 1]
         descricao_tipo_projeto = p.text[0:pos_numero]
         descricao_tipo_projeto = descricao_tipo_projeto.strip()
         achou = re.search('VETO.*AO', descricao_tipo_projeto)
         if achou:
            pos_numero = achou.span(0)[0] #posição string VETO
            pos_inicial = achou.span(0)[1]
            descricao_tipo_projeto=descricao_tipo_projeto[pos_inicial:len(descricao_tipo_projeto)].strip()

         if texto_numero_projeto != 'Sem Numero':
            id_projeto = texto_ano_projeto + dict_tipo_projeto[descricao_tipo_projeto][0] + id_projeto
         else:
            id_projeto = texto_numero_projeto + str(cont_projeto)

         #
         # --------------------------Identificando Vetos Parcial
         #
         achou = re.search("^.*VETO PARCIAL", p.text)
         if achou:
            flag_veto_parcial = 'Sim'
         #
         # --------------------------Identificando Vetos Total
         #
         achou = re.search("^.*VETO TOTAL", p.text)
         if achou:
            flag_veto_total = 'Sim'
         #
         # --------------------------Redação do Vencido e Redação Final
         #
         for item in lista_situacao_projeto_aux:
             achou = re.search("^.*REDAÇÃO DO VENCIDO", item)
             if achou:
                flag_redacao_vencido = 'Sim'
             achou = re.search("^.*REDAÇÃO FINAL", item)
             if achou:
                 flag_redacao_final = 'Sim'
         if texto_sessao != texto_sessao_anterior:
            texto_inclusao_ri =''
            numero_ordinal_projeto = 1
            texto_ordinal_projeto = str(numero_ordinal_projeto) + "º)"
            if flag_primeira_sessao == 'Sim':
               cont_sessao =1
               flag_primeira_sessao ='Não'
               #
               # --------------------------Criando a lista de atributos da sessão
               #
               lista_atributo_sessao_aux = []
               lista_atributo_sessao_aux.append(texto_sessao)
               lista_atributo_sessao_aux.append(texto_data_sessao)
               lista_atributo_sessao_aux.append(texto_dia_semana_sessao)
               lista_atributo_sessao_aux.append(texto_hora_sessao)
               dict_sessao.update({str(cont_sessao):lista_atributo_sessao_aux})
               lista_chave_projeto_aux=[]
               lista_chave_projeto_aux.append(id_projeto)
            else:
               dict_projeto_sessao.update({str(cont_sessao):lista_chave_projeto_aux})
               cont_sessao = cont_sessao + 1
               lista_atributo_sessao_aux = []
               lista_atributo_sessao_aux.append(texto_sessao)
               lista_atributo_sessao_aux.append(texto_data_sessao)
               lista_atributo_sessao_aux.append(texto_dia_semana_sessao)
               lista_atributo_sessao_aux.append(texto_hora_sessao)
               dict_sessao.update({str(cont_sessao): lista_atributo_sessao_aux})
               lista_chave_projeto_aux = []
               lista_chave_projeto_aux.append(id_projeto)
            texto_sessao_anterior = texto_sessao
            texto_hora_sessao = '' #Limpa hora da sessão anterior
         else:
            numero_ordinal_projeto = numero_ordinal_projeto + 1
            texto_ordinal_projeto = str(numero_ordinal_projeto) + "º)"
            lista_chave_projeto_aux.append(id_projeto)
            #if id_projeto not in lista_chave_projeto_aux:
            #  lista_chave_projeto_aux.append(id_projeto)

         lista_atributos_projeto_aux = []
         lista_atributos_projeto_aux.append(texto_numero_projeto)
         lista_atributos_projeto_aux.append(texto_ordinal_projeto)
         lista_atributos_projeto_aux.append(texto_inclusao_ri)
         lista_atributos_projeto_aux.append(lista_situacao_projeto_aux)
         lista_atributos_projeto_aux.append(texto_autores)
         lista_atributos_projeto_aux.append(texto_ementa)
         lista_atributos_projeto_aux.append(flag_veto_parcial)
         lista_atributos_projeto_aux.append(flag_veto_total)
         lista_atributos_projeto_aux.append(flag_redacao_vencido)
         lista_atributos_projeto_aux.append(flag_redacao_final)
         dict_projeto.update({id_projeto: lista_atributos_projeto_aux})
         continue
      #
      # --------------------------Identificando parágrafo PARECER(ES)
      #
      achou = re.search("^PARECER[ES]?:?|^PARECER[ES]? DA[S]? COMISS[ÃO|ÕES]:", p.text)
      if achou:
         #print(achou.span(0)[0])
         #pos_inicial = str.find(p.text, 'COMISSÃO')
         #texto_parecer = p.text[pos_inicial+9:].strip()
         texto_parecer = p.text
         #print(texto_parecer)
         dict_parecer.update({id_projeto:texto_parecer})
         continue
      #
      # --------------------------Identificando parágrafo NOVO(S) PARECER(ES)
      #
      achou = re.search("^[:space:]?NOVO[S]?.*PARECER[ES]:", p.text)
      if achou:
         pos_inicial = str.find(p.text, ':')
         texto_novo_parecer = p.text[pos_inicial+1:].strip()
         #texto_novo_parecer = p.text
         #print(texto_novo_parecer)
         dict_novo_parecer.update({id_projeto: texto_novo_parecer})
         continue

      #
      # --------------------------Identificando parágrafo NOVO(S) RELATOR(ES)
      #

      achou = re.search("^[:space:]?RELATOR[ES]?:", p.text)
      if achou:
         pos_inicial = str.find(p.text, ':')
         texto_relator = p.text[ pos_inicial+1:].strip()
         dict_relator.update({id_projeto:texto_relator})
         continue

      #
      # --------------------------Identificando parágrafo NOVO(S) RELATOR(ES)
      #
      achou = re.search("^.*PENDENDO.*DE.*PARECER[ES]?:|^.*PENDENDO.*DE.*PARECER[ES]? DA[S]? COMISS[ÃO|ÕES]:", p.text)
      if achou:
         pos_inicial = str.find(p.text, ':')
         pos_final = str.find(p.text, '.)')
         texto_comissao_pendente = p.text[pos_inicial+1:pos_final].strip()
         dict_comissao_pendente.update({id_projeto: texto_comissao_pendente})
         continue



   # -----------------------------Fim da Leitura do Documento de Origem (Documento da Alyne)
   #
   dict_projeto_sessao.update({str(cont_sessao):lista_chave_projeto_aux})

   for key in dict_projeto:
       lista_atributos_projeto_aux = dict_projeto[key]
       id_projeto = key # Número_chave do projeto
       url_busca=url_projeto+id_projeto
       texto_url = ''
       link_www3=''
       link_alerj=''
       #----------------------------------------------------------Busca Texto com tramitações de cada processo no ambiente Notes (tratar exceções)
       '''
       page = requests.get(url_busca)
       if page.status_code == 200:
          soup = BeautifulSoup(page.text, 'html.parser')
          texto_url = soup.prettify()
       link_alerj, link_www3 = busca_link_raiz(texto_url)
       '''

       # ---------------------------------------------------------

       lista_link_raiz = []
       lista_link_raiz.append(link_www3)
       lista_link_raiz.append(link_alerj)
       dict_link_raiz.update({id_projeto:lista_link_raiz})
       texto_autores = lista_atributos_projeto_aux[4]
       texto_ementa = lista_atributos_projeto_aux[5]
       flag_veto_parcial = lista_atributos_projeto_aux[6]
       flag_veto_total = lista_atributos_projeto_aux[7]
       flag_redacao_vencido = lista_atributos_projeto_aux[8]
       flag_redacao_final = lista_atributos_projeto_aux[9]

       texto_parecer = ''
       valor_projeto = dict_parecer.get(id_projeto, 'Não')
       if valor_projeto != 'Não':
          texto_parecer = dict_parecer[id_projeto]

       texto_novo_parecer = ''
       valor_projeto = dict_novo_parecer.get(id_projeto, 'Não')
       if valor_projeto != 'Não':
           texto_novo_parecer = dict_novo_parecer[id_projeto]

       texto_relator = ''
       valor_projeto = dict_relator.get(id_projeto, 'Não')
       if valor_projeto != 'Não':
           texto_relator = dict_relator[id_projeto]

       texto_comissao_pendente =''
       valor_projeto = dict_comissao_pendente.get(id_projeto, 'Não')
       if valor_projeto != 'Não':
           texto_comissao_pendente = dict_comissao_pendente[id_projeto]

       #------------------------------Busca palavras-chaves para veto parcial no Projeto
       #
       lista_chave_achada = []
       lista_chave_achada_aux=[]
       if flag_veto_parcial =='Sim':
          lista_chave_achada_aux = busca_palavras_chaves("VETO PARCIAL", dict_palavras_chaves,id_projeto,texto_url,dict_comissao_dominio)
          if len(lista_chave_achada_aux) >0:
              lista_chave_achada.append(lista_chave_achada_aux)
          lista_chave_achada_aux = busca_palavras_chaves("AUTÓGRAFO", dict_palavras_chaves,id_projeto,texto_url,dict_comissao_dominio)
          if len(lista_chave_achada_aux) > 0:
             lista_chave_achada.append(lista_chave_achada_aux)
       # ------------------------------Busca palavras-chaves para veto total da proposição
       #
       if flag_veto_total =='Sim':
          lista_chave_achada_aux = busca_palavras_chaves("VETO TOTAL", dict_palavras_chaves,id_projeto,texto_url,dict_comissao_dominio)
          if len(lista_chave_achada_aux) > 0:
             lista_chave_achada.append(lista_chave_achada_aux)
          lista_chave_achada_aux = busca_palavras_chaves("AUTÓGRAFO", dict_palavras_chaves,id_projeto,texto_url,dict_comissao_dominio)
          if len(lista_chave_achada_aux) > 0:
             lista_chave_achada.append(lista_chave_achada_aux)

       # ------------------------------Busca palavras-chaves para redação do vencido da proposição
       #
       if flag_redacao_vencido =='Sim':
          lista_chave_achada_aux = busca_palavras_chaves("REDAÇÃO DO VENCIDO", dict_palavras_chaves,id_projeto,texto_url,dict_comissao_dominio)
          if len(lista_chave_achada_aux) >0:
              lista_chave_achada.append(lista_chave_achada_aux)
       # ------------------------------Busca palavras-chaves para redação final da proposição
       #
       if flag_redacao_final =='Sim':
          lista_chave_achada_aux = busca_palavras_chaves("REDAÇÃO DO FINAL", dict_palavras_chaves,id_projeto,texto_url,dict_comissao_dominio)
          if len(lista_chave_achada_aux) > 0:
             lista_chave_achada.append(lista_chave_achada_aux)
       # ------------------------------Cria dicionário de busca de links da proposição
       #
       if len(lista_chave_achada)>0:
          dict_chave_projeto.update({id_projeto: lista_chave_achada})


       #------------------------------Busca palavras-chaves no texto de ementas
       #
       if len(texto_ementa) > 0:
          flag_busca_externa = 'Sim'
          lista_chave_achada = []
          lista_chave_achada = busca_palavras_chaves(texto_ementa, dict_palavras_chaves,id_projeto,texto_url,dict_comissao_dominio)
          if len(lista_chave_achada)>0:
             dict_chave_ementa.update({id_projeto: lista_chave_achada})
       #
       # ------------------------------Busca deputados(as) autores da matéria no texto de texto autores
       #
       if len(texto_autores) > 0:
           lista_autoria_achada = []
           lista_autoria_achada = busca_deputado(texto_autores)
           if len(lista_autoria_achada) > 0:
               dict_autores.update({id_projeto: lista_autoria_achada})
       #
       # ------------------------------Busca deputados(as) relatores da matéria no texto de texto relator
       #
       if len(texto_relator) > 0:
           lista_relatoria_achada = []
           lista_relatoria_achada = busca_deputado(texto_relator)
           if len(lista_relatoria_achada) > 0:
              dict_relator.update({id_projeto: lista_relatoria_achada})
       #
       #------------------------------Busca palavras-chaves no texto de pareceres
       #
       if len(texto_parecer) > 0:
          #flag_tipo_busca = 0
          lista_chave_achada = []
          lista_chave_achada = busca_palavras_chaves(texto_parecer, dict_palavras_chaves,id_projeto,texto_url,dict_comissao_dominio)
          if len(lista_chave_achada)>0:
             dict_chave_parecer.update({id_projeto: lista_chave_achada})
       #
       # ------------------------------Busca palavras-chaves no texto de novos pareceres
       #
       if len(texto_novo_parecer) > 0:
          #flag_tipo_busca = 1
          lista_chave_achada = []
          lista_chave_achada = busca_palavras_chaves(texto_novo_parecer, dict_palavras_chaves,id_projeto,texto_url,dict_comissao_dominio)
          if len(lista_chave_achada)> 0:
             dict_chave_novo_parecer.update({id_projeto: lista_chave_achada})
       #
       # ------------------------------Identifica as comissões com pendente de parecer no texto de comissao pendente
       #
       if len(texto_comissao_pendente)>0:
          #flag_tipo_busca = 1
          lista_chave_achada = []
          lista_chave_achada = busca_palavras_chaves(texto_comissao_pendente, dict_palavras_chaves,id_projeto,texto_url,dict_comissao_dominio)
          if len(lista_chave_achada)> 0:
             dict_chave_parecer_pendente.update({id_projeto: lista_chave_achada})
          lista_comissao_link = []
          lista_comissao_link = busca_link_comissao_pendente(texto_comissao_pendente,dict_comissao_dominio)
          if len(lista_comissao_link)> 0:
              dict_comissao_link.update({id_projeto:lista_comissao_link})

       #----------------------------------------------------------Fim Busca Texto com tramitações do processo no ambiente Notes (tratar exceções)

   return gera_pauta("pauta_final.docx",id_processamento,
              dict_sessao,
              dict_projeto_sessao,
              dict_projeto,
              dict_chave_projeto,
              dict_autores,
              dict_chave_ementa,
              dict_parecer,
              dict_novo_parecer,
              dict_relator,
              dict_comissao_link,
              dict_chave_parecer,
              dict_chave_novo_parecer,
              dict_chave_parecer_pendente,
              dict_tipo_projeto,
              dict_cores)

if __name__ == "__main__":
    import sys
    arquivo_entrada = sys.argv[1]
    print(main(arquivo_entrada,"00000"))




